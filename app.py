from flask import Flask, render_template, request, redirect, url_for, session, flash, make_response, jsonify
from flask_wtf.csrf import CSRFProtect
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_mail import Mail, Message
import sqlite3
import json
import os
import re
from werkzeug.security import check_password_hash, generate_password_hash
from functools import wraps
try:
    import cloudinary
    import cloudinary.uploader
    CLOUDINARY_AVAILABLE = True
except ImportError:
    CLOUDINARY_AVAILABLE = False


import os as _os

def _fix_broken_templates():
    """Fix any templates that have broken Jinja syntax from old injections."""
    import re as _re
    tmpl_dir = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "templates")
    
    for fname in ["edit_property.html", "edit_complex.html", "edit_unit.html"]:
        path = _os.path.join(tmpl_dir, fname)
        if not _os.path.exists(path):
            continue
        content = open(path, encoding="utf-8").read()
        lines = content.splitlines()
        # If file has suspicious bare Python-like assignments outside Jinja blocks
        # e.g. `prop.property_type=""` on its own line
        bad = [i+1 for i,l in enumerate(lines) 
               if _re.match(r"\s*(prop|unit)\.[a-z_]+=", l.strip()) 
               and "{%" not in l and "{{" not in l]
        if bad:
            # Remove those lines
            clean = [l for i,l in enumerate(lines) 
                     if (i+1) not in bad]
            open(path, "w", encoding="utf-8").write("\n".join(clean))

_fix_broken_templates()

app = Flask(__name__)
app.secret_key = "shivya_secret_key_2024"
app.config["WTF_CSRF_TIME_LIMIT"] = None
csrf = CSRFProtect(app)
@app.errorhandler(500)
def internal_error(e):
    import traceback
    return "<pre style='padding:20px;background:#111;color:#f87171;font-size:13px;'>" + traceback.format_exc() + "</pre>", 500


limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=[],
    storage_uri="memory://"
)

# -----------------------------------------------
# EMAIL CONFIG — change these to your Gmail details
# -----------------------------------------------
app.config["MAIL_SERVER"]   = "smtp.hostinger.com"
app.config["MAIL_PORT"]     = 465
app.config["MAIL_USE_TLS"]  = False
app.config["MAIL_USE_SSL"]  = True
app.config["MAIL_USERNAME"] = "hello@shivyaspaces.com"
app.config["MAIL_PASSWORD"] = "Shivya@768"
app.config["MAIL_DEFAULT_SENDER"] = ("ShivyaSpaces", "hello@shivyaspaces.com")

# -----------------------------------------------
# SITE URL — change this after hosting
# -----------------------------------------------
SITE_URL = os.environ.get("SITE_URL", "https://www.shivyaspaces.com")
mail = Mail(app)

# ── Cloudinary config (optional - set env vars to enable) ──
if CLOUDINARY_AVAILABLE:
    import cloudinary
    import cloudinary.uploader
    cloudinary.config(
        cloud_name = os.environ.get("CLOUDINARY_CLOUD_NAME", ""),
        api_key    = os.environ.get("CLOUDINARY_API_KEY", ""),
        api_secret = os.environ.get("CLOUDINARY_API_SECRET", ""),
        secure     = True
    )

# ── Upload folder for local image storage (fallback) ──
# Use /tmp for Railway (persistent static dir may not be writable)
_static_uploads = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "uploads")
_tmp_uploads = "/tmp/shivyaspaces_uploads"
try:
    os.makedirs(_static_uploads, exist_ok=True)
    # Test write
    _test = os.path.join(_static_uploads, ".test")
    open(_test, "w").close()
    os.remove(_test)
    UPLOAD_FOLDER = _static_uploads
    UPLOAD_URL_BASE = "/static/uploads"
except Exception:
    os.makedirs(_tmp_uploads, exist_ok=True)
    UPLOAD_FOLDER = _tmp_uploads
    UPLOAD_URL_BASE = "/tmp-uploads"
ALLOWED_EXTENSIONS = {"jpg", "jpeg", "png", "webp", "gif", "bmp", "tiff", "tif", "heic", "heif", "avif", "svg"}
MAX_UPLOAD_SIZE = 10 * 1024 * 1024  # 10MB


# -----------------------------------------------
# HELPERS
# -----------------------------------------------

def get_db():
    conn = sqlite3.connect("properties.db", timeout=10)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def get_setting(key, default=""):
    """Read a setting from the database."""
    try:
        conn = sqlite3.connect("properties.db", timeout=10)
        conn.row_factory = sqlite3.Row
        conn.execute("""CREATE TABLE IF NOT EXISTS settings
            (key TEXT PRIMARY KEY, value TEXT)""")
        row = conn.execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
        conn.close()
        return row["value"] if row else default
    except Exception:
        return default


def save_setting(key, value):
    """Save a setting to the database."""
    try:
        conn = sqlite3.connect("properties.db", timeout=10)
        conn.execute("""CREATE TABLE IF NOT EXISTS settings
            (key TEXT PRIMARY KEY, value TEXT)""")
        conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?,?)", (key, value))
        conn.commit()
        conn.close()
        return True
    except Exception:
        return False


def get_mail_config():
    """Get current mail config from DB, fallback to hardcoded defaults."""
    return {
        "username": get_setting("mail_username", app.config["MAIL_USERNAME"]),
        "password": get_setting("mail_password", app.config["MAIL_PASSWORD"]),
        "sender_name": get_setting("mail_sender_name", "ShivyaSpaces"),
    }


def send_welcome_email(owner_name, owner_email, username, password, slug):
    """Send welcome email using mail config from database."""
    try:
        cfg = get_mail_config()
        # Update mail config live from DB
        app.config["MAIL_USERNAME"]       = cfg["username"]
        app.config["MAIL_PASSWORD"]       = cfg["password"]
        app.config["MAIL_DEFAULT_SENDER"] = (cfg["sender_name"], cfg["username"])
        mail.init_app(app)

        msg = Message(
            subject="Your ShivyaSpaces Account is Ready",
            recipients=[owner_email]
        )
        msg.html = f"""
        <div style="font-family:Inter,Arial,sans-serif;max-width:520px;margin:0 auto;padding:32px 24px;color:#111827;">
            <h2 style="font-size:22px;font-weight:800;margin-bottom:4px;">Welcome to ShivyaSpaces</h2>
            <p style="color:#6b7280;font-size:14px;margin-bottom:28px;">Your account has been created. Here are your login credentials.</p>
            <div style="background:#f8fafc;border:1px solid #e5e7eb;border-radius:14px;padding:20px;margin-bottom:24px;">
                <p style="font-size:12px;font-weight:600;color:#9ca3af;text-transform:uppercase;letter-spacing:0.5px;margin-bottom:16px;">Your Credentials</p>
                <table style="width:100%;border-collapse:collapse;">
                    <tr><td style="font-size:13px;color:#6b7280;padding:8px 0;width:40%;">Login URL</td><td style="font-size:13px;font-weight:600;color:#111827;">/login</td></tr>
                    <tr><td style="font-size:13px;color:#6b7280;padding:8px 0;">Email</td><td style="font-size:13px;font-weight:600;color:#111827;">{owner_email}</td></tr>
                    <tr><td style="font-size:13px;color:#6b7280;padding:8px 0;">Password</td><td style="font-size:13px;font-weight:600;color:#111827;">{password}</td></tr>
                    <tr><td style="font-size:13px;color:#6b7280;padding:8px 0;">Your Page</td><td style="font-size:13px;font-weight:600;color:#111827;">/{slug}</td></tr>
                </table>
            </div>
            <p style="font-size:13px;color:#6b7280;line-height:1.7;">Log in to your dashboard to start adding your properties. Your public page will be live as soon as you add your first listing.</p>
            <hr style="border:none;border-top:1px solid #e5e7eb;margin:24px 0;" />
            <p style="font-size:12px;color:#9ca3af;">This email was sent by {cfg["sender_name"]}. Please keep your credentials safe.</p>
        </div>
        """
        mail.send(msg)
        return True
    except Exception as e:
        print(f"Email send failed: {e}")
        return False


def parse_property(row):
    if row is None:
        return None
    p = dict(row)
    p["amenities"]      = json.loads(p["amenities"])      if p.get("amenities")      else []
    p["gallery_images"] = json.loads(p["gallery_images"]) if p.get("gallery_images") else []
    # Safe defaults for optional fields
    for field in ["title","location","rent","status","description","property_type",
                  "facing","maintenance","floor","built","security_deposit","image",
                  "visibility","listing_type","sale_price"]:
        if field not in p or p[field] is None:
            p[field] = ""
    for field in ["beds","baths","sqft","is_complex","parent_id","contact_for_price"]:
        if field not in p or p[field] is None:
            p[field] = 0
    # defaults
    if not p["visibility"]:   p["visibility"]   = "public"
    if not p["listing_type"]: p["listing_type"] = "rent"
    return p



def migrate_db():
    """Add new columns if they don't exist."""
    conn = get_db()
    existing = [r[1] for r in conn.execute("PRAGMA table_info(properties)").fetchall()]
    migrations = [
        ("visibility",    "TEXT DEFAULT 'public'"),
        ("listing_type",  "TEXT DEFAULT 'rent'"),
        ("sale_price",    "TEXT DEFAULT ''"),
        ("contact_for_price", "INTEGER DEFAULT 0"),
    ]
    for col, definition in migrations:
        if col not in existing:
            conn.execute(f"ALTER TABLE properties ADD COLUMN {col} {definition}")
            print(f"Migration: added column {col}")

    # Create leads table if not exists
    conn.execute("""CREATE TABLE IF NOT EXISTS leads (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        property_id INTEGER NOT NULL,
        tenant_name TEXT NOT NULL,
        tenant_phone TEXT NOT NULL,
        status TEXT DEFAULT 'new',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (property_id) REFERENCES properties(id)
    )""")
    # Add status column if missing (existing installs)
    lead_cols = [r[1] for r in conn.execute("PRAGMA table_info(leads)").fetchall()]
    if "status" not in lead_cols:
        conn.execute("ALTER TABLE leads ADD COLUMN status TEXT DEFAULT 'new'")
        print("Migration: added leads.status")

    # Create rental_agreements table
    conn.execute("""CREATE TABLE IF NOT EXISTS rental_agreements (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        lead_id INTEGER NOT NULL,
        property_id INTEGER NOT NULL,
        owner_id INTEGER NOT NULL,
        executed_date TEXT,
        effective_date TEXT,
        lockin_end_date TEXT,
        owner_age TEXT,
        owner_address TEXT,
        tenant_name TEXT,
        tenant_mobile TEXT,
        tenant_aadhaar TEXT,
        tenant_pan TEXT,
        tenant_father TEXT,
        tenant_age TEXT,
        tenant_address TEXT,
        prop_flat TEXT,
        prop_floor TEXT,
        prop_facing_type TEXT,
        prop_building TEXT,
        prop_street TEXT,
        prop_area TEXT,
        prop_taluk TEXT,
        prop_pin TEXT,
        prop_state TEXT,
        monthly_rent TEXT,
        monthly_rent_revised TEXT,
        rent_due_day TEXT,
        security_deposit TEXT,
        lockin_months TEXT,
        enhancement_pct TEXT,
        notice_period TEXT,
        hall_foyer_light TEXT, hall_tube_light TEXT, hall_led_light TEXT,
        hall_curtain_rod TEXT, hall_wall_bulb TEXT, hall_fan TEXT,
        hall_tv_cabinet TEXT, hall_ups TEXT, hall_bell TEXT, hall_pop_lights TEXT,
        kitchen_modular TEXT, kitchen_led TEXT, kitchen_led_utility TEXT,
        pooja_structure TEXT, pooja_storage TEXT, pooja_bulb TEXT,
        bed_fan TEXT, bed_tube TEXT, bed_curtain TEXT, bed_loft TEXT,
        bed_sliding_lock TEXT, bed_drawer_locker TEXT, bed_dressing TEXT,
        bed_bangles_rod TEXT, bed_hanger TEXT, bed_bulb TEXT,
        bed_ceiling TEXT, bed_desk TEXT,
        bath_storage TEXT, bath_geyser TEXT, bath_bulb TEXT,
        bath_ceiling TEXT, bath_exhaust TEXT,
        key_main_num TEXT, key_main_qty TEXT,
        key_bedroom_num TEXT, key_bedroom_qty TEXT,
        key_sliding_num TEXT, key_sliding_qty TEXT,
        key_drawer_num TEXT, key_drawer_qty TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    conn.commit()
    conn.close()

migrate_db()

def clean_phone(phone):
    """
    Returns digits only for wa.me links e.g. 919876543210
    Strips +, spaces, dashes — everything except digits.
    """
    if not phone:
        return ""
    phone = phone.strip()
    # Strip everything except digits
    digits = re.sub(r"[^\d]", "", phone)
    # Handle 00-prefix international numbers
    if digits.startswith("00"):
        digits = digits[2:]
    # Bare 10 digits = assume India
    if len(digits) == 10:
        digits = "91" + digits
    return digits


def format_phone_display(phone):
    """
    Returns phone for display with + prefix e.g. +91 98765 43210
    """
    if not phone:
        return ""
    digits = clean_phone(phone)
    if digits:
        return "+" + digits
    return phone


def validate_phone(phone):
    """
    Accepts international formats:
    - With country code: +91 98765 43210, +1 415 555 0132 etc.
    - Without country code: 9876543210 (assumed India)
    - Must have at least 10 digits after stripping formatting
    - Must have at most 15 digits (international standard E.164)
    """
    if not phone:
        return False, "Phone number is required."
    digits = re.sub(r"[^\d]", "", phone)
    if len(digits) < 10:
        return False, "Phone number must have at least 10 digits."
    if len(digits) > 15:
        return False, "Phone number is too long. Maximum 15 digits."
    return True, ""


def validate_email(email):
    """
    Checks:
    - Has exactly one @
    - Has a domain with at least one dot after @
    - No spaces
    - Local part (before @) is not empty
    - Domain part has valid characters
    """
    if not email:
        return False, "Email address is required."
    if " " in email:
        return False, "Email address cannot contain spaces."
    pattern = r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$'
    if not re.match(pattern, email):
        return False, "Please enter a valid email address (e.g. name@example.com)."
    return True, ""


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "owner_id" not in session:
            flash("Please log in to access that page.", "warning")
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "admin_id" not in session:
            flash("Admin access required.", "warning")
            return redirect(url_for("admin_login"))
        return f(*args, **kwargs)
    return decorated


def nocache(resp):
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"]        = "no-cache"
    resp.headers["Expires"]       = "0"
    return resp


# -----------------------------------------------
# 404 HANDLER
# -----------------------------------------------




@app.route("/rental-agreement/<int:lead_id>", methods=["GET"])
@login_required
def rental_agreement_form(lead_id):
    """Show pre-filled rental agreement form for a lead."""
    conn = get_db()

    lead = conn.execute("SELECT * FROM leads WHERE id=?", (lead_id,)).fetchone()
    if not lead:
        flash("Lead not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))

    prop = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=?",
        (lead["property_id"], session["owner_id"])
    ).fetchone())

    if not prop:
        flash("Property not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))

    owner = conn.execute("SELECT * FROM owners WHERE id=?", (session["owner_id"],)).fetchone()

    # Check for existing saved agreement
    existing = conn.execute(
        "SELECT * FROM rental_agreements WHERE lead_id=? ORDER BY id DESC LIMIT 1",
        (lead_id,)
    ).fetchone()
    conn.close()

    existing_data = dict(existing) if existing else {}

    return render_template("rental_agreement_form.html",
                           lead=dict(lead),
                           prop=prop,
                           owner=dict(owner),
                           existing=existing_data,
                           brand_name=session["brand_name"],
                           owner_name=session["owner_name"])


@app.route("/rental-agreement/<int:lead_id>/generate", methods=["POST"])
@login_required
def generate_rental_agreement(lead_id):
    """Save form data and generate Word document."""
    import os

    conn = get_db()
    lead = conn.execute("SELECT * FROM leads WHERE id=?", (lead_id,)).fetchone()
    if not lead:
        flash("Lead not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))

    prop = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=?",
        (lead["property_id"], session["owner_id"])
    ).fetchone())
    owner = conn.execute("SELECT * FROM owners WHERE id=?", (session["owner_id"],)).fetchone()

    # Collect all form fields
    f = request.form
    data = {
        "lead_id": lead_id,
        "property_id": lead["property_id"],
        "owner_id": session["owner_id"],
        "executed_date":     f.get("executed_date",""),
        "effective_date":    f.get("effective_date",""),
        "lockin_end_date":   f.get("lockin_end_date",""),
        "owner_age":         f.get("owner_age",""),
        "owner_address":     f.get("owner_address",""),
        "tenant_name":       f.get("tenant_name",""),
        "tenant_mobile":     f.get("tenant_mobile",""),
        "tenant_aadhaar":    f.get("tenant_aadhaar",""),
        "tenant_pan":        f.get("tenant_pan",""),
        "tenant_father":     f.get("tenant_father",""),
        "tenant_age":        f.get("tenant_age",""),
        "tenant_address":    f.get("tenant_address",""),
        "prop_flat":         f.get("prop_flat",""),
        "prop_floor":        f.get("prop_floor",""),
        "prop_facing_type":  f.get("prop_facing_type",""),
        "prop_building":     f.get("prop_building",""),
        "prop_street":       f.get("prop_street",""),
        "prop_area":         f.get("prop_area",""),
        "prop_taluk":        f.get("prop_taluk",""),
        "prop_pin":          f.get("prop_pin",""),
        "prop_state":        f.get("prop_state",""),
        "monthly_rent":      f.get("monthly_rent",""),
        "monthly_rent_revised": f.get("monthly_rent_revised",""),
        "rent_due_day":      f.get("rent_due_day",""),
        "security_deposit":  f.get("security_deposit",""),
        "lockin_months":     f.get("lockin_months",""),
        "enhancement_pct":   f.get("enhancement_pct",""),
        "notice_period":     f.get("notice_period",""),
        # Fittings
        "hall_foyer_light":  f.get("hall_foyer_light",""),
        "hall_tube_light":   f.get("hall_tube_light",""),
        "hall_led_light":    f.get("hall_led_light",""),
        "hall_curtain_rod":  f.get("hall_curtain_rod",""),
        "hall_wall_bulb":    f.get("hall_wall_bulb",""),
        "hall_fan":          f.get("hall_fan",""),
        "hall_tv_cabinet":   f.get("hall_tv_cabinet",""),
        "hall_ups":          f.get("hall_ups",""),
        "hall_bell":         f.get("hall_bell",""),
        "hall_pop_lights":   f.get("hall_pop_lights",""),
        "kitchen_modular":   f.get("kitchen_modular",""),
        "kitchen_led":       f.get("kitchen_led",""),
        "kitchen_led_utility": f.get("kitchen_led_utility",""),
        "pooja_structure":   f.get("pooja_structure",""),
        "pooja_storage":     f.get("pooja_storage",""),
        "pooja_bulb":        f.get("pooja_bulb",""),
        "bed_fan":           f.get("bed_fan",""),
        "bed_tube":          f.get("bed_tube",""),
        "bed_curtain":       f.get("bed_curtain",""),
        "bed_loft":          f.get("bed_loft",""),
        "bed_sliding_lock":  f.get("bed_sliding_lock",""),
        "bed_drawer_locker": f.get("bed_drawer_locker",""),
        "bed_dressing":      f.get("bed_dressing",""),
        "bed_bangles_rod":   f.get("bed_bangles_rod",""),
        "bed_hanger":        f.get("bed_hanger",""),
        "bed_bulb":          f.get("bed_bulb",""),
        "bed_ceiling":       f.get("bed_ceiling",""),
        "bed_desk":          f.get("bed_desk",""),
        "bath_storage":      f.get("bath_storage",""),
        "bath_geyser":       f.get("bath_geyser",""),
        "bath_bulb":         f.get("bath_bulb",""),
        "bath_ceiling":      f.get("bath_ceiling",""),
        "bath_exhaust":      f.get("bath_exhaust",""),
        "key_main_num":      f.get("key_main_num",""),
        "key_main_qty":      f.get("key_main_qty",""),
        "key_bedroom_num":   f.get("key_bedroom_num",""),
        "key_bedroom_qty":   f.get("key_bedroom_qty",""),
        "key_sliding_num":   f.get("key_sliding_num",""),
        "key_sliding_qty":   f.get("key_sliding_qty",""),
        "key_drawer_num":    f.get("key_drawer_num",""),
        "key_drawer_qty":    f.get("key_drawer_qty",""),
    }

    # Save to DB (upsert)
    existing = conn.execute(
        "SELECT id FROM rental_agreements WHERE lead_id=?", (lead_id,)
    ).fetchone()

    cols = ", ".join(data.keys())
    placeholders = ", ".join(["?"] * len(data))
    vals = list(data.values())

    if existing:
        set_clause = ", ".join([f"{k}=?" for k in data.keys() if k not in ("lead_id","property_id","owner_id")])
        update_vals = [data[k] for k in data.keys() if k not in ("lead_id","property_id","owner_id")]
        conn.execute(f"UPDATE rental_agreements SET {set_clause} WHERE lead_id=?", update_vals + [lead_id])
    else:
        conn.execute(f"INSERT INTO rental_agreements ({cols}) VALUES ({placeholders})", vals)
    conn.commit()

    conn.close()

    # Validate required fields
    required_fields = {
        "executed_date": "Executed Date",
        "effective_date": "Effective Date",
        "lockin_end_date": "Lock-in End Date",
        "owner_age": "Owner Age",
        "owner_address": "Owner Address",
        "tenant_name": "Tenant Name",
        "tenant_mobile": "Tenant Mobile",
        "tenant_aadhaar": "Aadhaar Number",
        "tenant_pan": "PAN Number",
        "tenant_father": "Father's Name",
        "tenant_age": "Tenant Age",
        "tenant_address": "Tenant Address",
        "prop_flat": "Flat Number",
        "prop_floor": "Floor",
        "prop_building": "Building Name",
        "prop_street": "Street",
        "prop_pin": "PIN Code",
        "monthly_rent": "Monthly Rent",
        "security_deposit": "Security Deposit",
    }
    missing = [label for key, label in required_fields.items() if not data.get(key, "").strip()]
    if missing:
        flash(f"Please fill in all required fields: {', '.join(missing)}", "danger")
        return redirect(url_for("rental_agreement_form", lead_id=lead_id))

    # Build data dict with owner details
    doc_data = dict(data)
    doc_data["owner_name"]   = owner["name"]
    doc_data["owner_mobile"] = owner["phone"] or ""

    # Generate Word doc using python-docx
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        d = doc_data

        def get(key, default=""):
            return d.get(key) or default

        lock_months = get("lockin_months", "6")
        lock_word   = {"6":"Six","11":"Eleven"}.get(lock_months, lock_months)
        enh_pct     = get("enhancement_pct", "8")
        enh_word    = {"8":"Eight","10":"Ten","5":"Five"}.get(enh_pct, enh_pct)
        rent_day    = get("rent_due_day", "5th")
        notice      = get("notice_period", "30 days")
        prop_full   = (
            f"{get('prop_flat')}, {get('prop_floor')}, "
            + (f"{get('prop_facing_type')}, " if get('prop_facing_type') else "")
            + f"{get('prop_building')}, Site #37, Ambica Arcade Phase-1, "
            + f"{get('prop_street')}, near Brigade Orchid North gate, Devanahalli, "
            + f"Bangalore-{get('prop_pin')}, {get('prop_state')}"
        )

        doc = DocxDocument()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(0.75)

        def set_font(run, bold=False, size=11):
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(size)
            run.font.bold  = bold

        def add_para(text_parts, align="justify"):
            """text_parts: list of (str, bool) tuples or plain strings."""
            para = doc.add_paragraph()
            if align == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "justify":
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Flatten: if text_parts is a list containing a single list, unwrap it
            if len(text_parts) == 1 and isinstance(text_parts[0], list):
                text_parts = text_parts[0]
            for part in text_parts:
                if isinstance(part, tuple):
                    txt, bold = part
                elif isinstance(part, str):
                    txt, bold = part, False
                else:
                    txt, bold = str(part), False
                run = para.add_run(str(txt))
                set_font(run, bold=bool(bold))
            return para

        def heading(text, size=12):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            set_font(run, bold=True, size=size)
            return para

        def clause(number, title):
            para = doc.add_paragraph()
            run = para.add_run(f"{number}. {title}:")
            set_font(run, bold=True)
            return para

        def body(parts_or_text, align="justify"):
            """Accept either a list of (str,bool) tuples or a plain string."""
            if isinstance(parts_or_text, str):
                return add_para([(parts_or_text, False)], align)
            return add_para([parts_or_text] if not isinstance(parts_or_text, list) else parts_or_text, align)

        def sp():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

        # ── TITLE ──
        heading("RENTAL AGREEMENT", 14)
        sp()

        body([
            "This ", False,
            ("RENTAL AGREEMENT", True),
            (", is made and executed on this ", False),
            (get("executed_date"), True),
            (' ("Agreement") at Bangalore and effective from ', False),
            (get("effective_date"), True),
            (" by and", False),
        ])
        sp()
        add_para([("BETWEEN", True)], "center")
        sp()

        # Owner
        body([
            (f"Mr. {get('owner_name')}", True),
            (f", Age: {get('owner_age')}, Residing at: {get('owner_address')}. Mobile: ", False),
            (get("owner_mobile"), True),
            (".", False),
        ])
        sp()
        body([
            ('HEREINAFTER called the ', False),
            ('"LESSOR / OWNER / Landlord"', True),
            (' which expression shall unless repugnant to the meaning for content thereof, mean and include his executors, legal representatives, administrators, successors in title and assigns of the ', False),
            ('ONE PART.', True),
        ])
        sp()
        body([("AND", True)], )
        sp()

        # Tenant
        body([
            (f"Mr. {get('tenant_name')}", True),
            (f" (Aadhar No: {get('tenant_aadhaar')}) and PAN No: ({get('tenant_pan')}), S/O: {get('tenant_father')}, {get('tenant_address')}. Mobile: ", False),
            (get("tenant_mobile"), True),
        ])
        sp()
        body([
            ('HEREINAFTER called the ', False),
            ('"LESSEE / TENANT"', True),
            (' which expression shall unless repugnant to the meaning for content thereof, mean and include his executors, legal representatives, administrators, successors in title and assigns of the ', False),
            ('OTHER/SECOND PART.', True),
        ])
        sp()

        # WHEREAS
        body([
            ('WHEREAS the Lessor is the absolute owner of property bearing: ', False),
            (prop_full, True),
            (', Karnataka. More fully described in the Schedule given hereunder and hereinafter referred to as the Schedule Property.', False),
        ])
        sp()
        body([("AND WHEREAS the Lessee is in need of accommodation for his residential purpose and the Lessee has approached the Lessor for the grant of the lease of the Schedule Property, and the Lessor has agreed for the same.", False)])
        sp()
        body([("Which is more fully described in the schedule herein under and hereinafter referred to as the schedule property/flat and whereas accordingly this agreement witness that the Lessee hereby agreed to occupy the schedule property/flat on the terms and conditions hereinafter mentioned.", False)])
        sp()
        body([("NOW THEREFORE IT IS HEREBY AGREED TO, DECLARED AND RECORDED BY AND BETWEEN THE PARTIES HERETO THIS RENTAL AGREEMENT AND WITNESS AS FOLLOWS", True)])
        sp()

        # CLAUSE 1
        clause(1, "Duration")
        body([("That the Lessor hereby allow the Lessee herein a revocable leave and license, to occupy the licensed premises, described in the scheduled property hereunder written without creating any tenancy rights or any other grants/rights, titles and interest in favour of Lessee.", False)])
        sp()
        body([
            ("The duration of the lease will be initially for a period of maximum 11 months commencing from ", False),
            (get("effective_date"), True),
            (f" and is subject to renewal for a further period of 11 months with an increase in rent of {enh_pct}.0% on mutually agreeable and renewable according to the mutual understanding of both the parties by executing a fresh Rental Agreement, fulfilling statutory requirement and the expenses shall be met by the LESSEE terms & conditions.", False),
        ])
        sp()

        # CLAUSE 2
        clause(2, "Rent")
        body([
            ("The monthly advance rent (pay & stay) payable by the Lessee to the Lessor for the schedule property shall be ", False),
            (f"Rs. {get('monthly_rent')} ", True),
            ("as a rent per month, excluding of electricity, internet, cable T.V fee and gas charges. Incase if the owner's association/ Community/builder takes a decision in the future considering the market increase and labor demand on the maintenance charges then LESSOR will inform the same to LESSE on the same who (LESSEE) needs to pay the maintenance charges per month accordingly.", False),
        ])
        sp()
        body([(f"The advance rent shall be payable by Lessee before {rent_day} of every English calendar month to be credited directly to the LESSOR's following account. No exceptions/exemptions are allowed considering any other circumstances like natural climates, COVID-19 or any other Govt. imposed rules until agreed with Lessor/Owner.", False)])
        sp()
        body([("If the Lessee delays payment of the advance Rent beyond the timeline specified for any month, the Lessor shall be entitled to receive an interest at 18% per annum for the period of delay days.", False)])
        sp()
        body([
            ("The electricity arrangement (more explained in Temporary Common Electricity Supply & Cost Sharing section) is temporary until a permanent dedicated direct meter connection is installed/commenced. Upon activation of the direct meter connection, the monthly rent will be auto revised to ", False),
            (f"₹{get('monthly_rent_revised')} ", True),
            ("and the tenant shall pay the electricity charges directly to BESCOM based on actual consumption.", False),
        ])
        sp()

        # CLAUSE 3
        clause(3, "Lock-in Period")
        body([
            (f"The Parties agree that neither Party shall be entitled to terminate this Agreement for a period of {lock_months} ({lock_word}) months starting with effect from the Commencement Date (i.e., from ", False),
            (get("effective_date"), True),
            (" till ", False),
            (get("lockin_end_date"), True),
            (') ("Lock-in Period"). In the event, the Lessee terminates the Agreement before the expiry of the Lock-in Period, the Lessor shall be entitled to deduct the Security Deposit in addition to the rent payable for the Notice Period as follows:', False),
        ])
        sp()
        body([("(a) If the termination date falls within one month from the lock in expiry date, an amount equivalent to one month's rent shall be deducted from the Security Deposit", False)])
        body([("(b) If the termination date exceeds one month from the rental expiry date, an amount equivalent to two months' rent shall be deducted from the Security Deposit.", False)])
        sp()

        # CLAUSE 4
        clause(4, "Enhancement")
        body([(f"The rent payable to the LESSOR by the LESSEE as aforesaid in Clause No.1 shall be enhanced by {enh_pct}% ({enh_word} Percent) of the last amount paid at the end of 11 months, subject to the renewal of the rental agreement by mutual consent of both the parties.", False)])
        sp()

        # CLAUSE 5
        clause(5, "Interest Free Security Deposit")
        body([
            ("An interest free security deposit of ", False),
            (f"Rs. {get('security_deposit')} ", True),
            ("by way of account transfer from Lessee to Lessor account online. The same will be refundable by the lessor to the lessee at the time of vacating the scheduled property (termination of tenancy agreement). The possession of the property is deemed complete only upon handover of all keys. The deposit shall be refunded within 7 days of vacating the premises & hand over the keys after adjustment all dues.", False),
        ])
        sp()
        body([("To the LESSOR as security deposit, which the LESSOR hereby acknowledges the said amount and the same shall be held by the LESSOR as Security Deposit during the continuance of the tenancy and/or any extension thereof and shall be repaid to the LESSEE free of interest at the end of the period of the lease, at the time of LESSEE vacating and delivering the said premises in the same condition, in which it was let out, or on termination of this agreement and after deducting the dues payable (like rent, maintenance, electricity, internet, damages and other arrears), if any, by the LESSEE to the LESSOR, painting/deep cleaning of the house and cost of rectification of any damages done.", False)])
        sp()

        # CLAUSES 6-27 — static text with variable substitution
        for num, title, text in [
            (6, "Temporary Common Electricity Supply & Cost Sharing Temporary Electricity Arrangement",
             "The Lessee is aware that, as on the date of execution of this Agreement, the building is supplied with a temporary sanctioned common electrical load of 1 KW, which is shared by all occupants of the building. Separate permanent electricity meter connections for individual flats have been applied for with BESCOM and are awaited.||PARA||Towards the common electricity expenses, the Lessee agrees to pay a fixed monthly contribution of Rs. 1,000/- (Rupees One Thousand only) along with the monthly rent/license fee, until individual permanent meters are installed.||PARA||Upon installation/commission of dedicated permanent electricity meter(s) for the Schedule Property (individual flats) by BESCOM, this temporary electricity arrangement shall automatically stand terminated, and electricity charges shall thereafter be paid directly by the Lessee as per the individual consumption usage, meter readings and bills."),
            (7, "Gas / Electricity / Internet / Cable T.V",
             "The Lessee shall maintain the schedule property in a state of good housekeeping and condition. The LESSEE shall bear and pay the charges for the Electricity bill, Internet bill & Gas bill for the premised shall be borne by the Lessee based on the usage/fixed by the Owner welfare association / Builder (if association not yet formed) / respective authorities/suppliers on every month along with Rentals."),
            (8, "Property Tax",
             "The property tax and all other taxes, rates, cesses, assessments, duties and other outgoing payable related to municipal Corporation/ Grama Panchayat or any other authority of the government shall be borne by the LESSOR."),
            (9, "Internal Maintenance",
             "The Lessee shall maintain the schedule property in a state of good order and condition and shall not cause any damage to disfigurement to the schedule property or to any wall paintings, wardrobes, kitchen wood works, bath room, fittings and fixtures (as listed at the end of this contract) therein always. Any damage caused by the Lessee shall be made good by the Lessee or an equivalent amount will be deducted from the security deposit at the time of vacating the scheduled property.||PARA||Repairs And Maintenance: It is specifically agreed between the parties to this deed that the cost of repairs and maintenance of the furniture, fixtures, electrical fittings like bulbs, tube lights, fans, fan regulators, etc. and water fittings like taps, shower, faucets, cistern, etc. shall be borne by the Lessee alone. However, problems related to civil and construction work, concealed electrical wiring and concealed plumbing, seepage and leakage of walls, ceiling or any part of the Schedule Property, shall be handled and set right by the Lessors at their own cost. However, the Lessors shall ensure that electrical fittings like tube lights, bulbs, fans and fan regulators, call bell etc. and sanitary fixtures like taps, cisterns, showers, geysers etc. are in perfectly working condition, at the time of handing over the schedule property to the Lessee, failing which, the Lessee is at liberty to get the same repaired or replaced and recover this cost from the Lessors. At the time of vacating the Schedule Property and handing it back to the Lessors, the Lessee shall ensure that all electrical and sanitary fittings are in the same working condition, The Lessees will ensure that the Schedule Property is handed over back to the Lessors in neat and tidy condition and all trash like bottles, cardboard and corrugated boxes, old newspapers, batteries adhesive packing tapes, used bulbs, tube lights, buckets, plastic mugs, toilet brushes, scrubbers, detergents and any type of junk is completely cleared before vacating the Schedule property. Excessive nailing in the Schedule Property and disfiguring of any part of the Schedule Property IN ANY MANNER, thus rendering it non-tenantable, is strictly inadmissible, failing which, the Lessee shall bear the entire cost of repairing to restore THIS damage caused, to its perfect condition. The Lessee shall not cause damage to the Schedule Property during his occupation of the Schedule Property. After ensuring that there are no damages or arrears OR after adjusting the cost of the damages and/or arrears, the balance or full amount (AS THE CASE MAY BE) will be refunded by the Lessor to the Lessee immediately the Lessee vacates the Schedule Property."),
            (10, "Additions & Alterations",
             "The Lessee shall not be entitled to make any additions or alteration internal and/or external to the said Scheduled premises and to the furniture, fixtures and electrical fittings installed/provided by the LESSOR at the Scheduled Premises which involves structural changes or not. The Lessee shall however be entitled to fix telephones, televisions, etc in the provision already given in the scheduled property. The Lessee shall be entitled to remove the same at the time of vacating the schedule property and shall be made good to the full satisfaction of the Lessor. The LESSEE shall ensure that the damages to walls and fixtures, if any, are duly repaired before handing over the possession."),
            (11, "Nature of Usage / Purpose permitted",
             "The Lessee shall use the schedule property only for Residential Purpose and that too with a family. The Lessee shall not keep or store in or upon any part of the schedule premises any goods of combustible or explosive nature, except those for cooking purposes and shall not store any offensive items, which may cause damage to the house premises and/or shall not carry any business activities either legal or illegal or un-lawful and illegitimate activities. Any kind of disturbance, misbehavior, unethical practice may call for immediate termination of agreement and evacuation from the property with a penalty of three months rental deducted from the advance security amount."),
            (12, "No Tenancy",
             "The LESSEE will not have any right to transfer, assign, and sub-let or grant any license or sub-license in respect of the scheduled property and premises or any part thereof and also shall not mortgage or raise any loan against the said premises. The LESSEE shall use the Scheduled Premises in a reasonable manner without causing any disturbance to the neighbors and agree to abide by the rules and regulations of Ambika Arcade Owners Association. That, the Lessee shall not claim any tenancy rights."),
            (13, "Inspection",
             "The Lessee shall permit the Lessor or his/her designated representatives, during the reasonable hours in the day time and upon making prior (at least 2 days in advance) appointment with Lessee to visit / inspect the schedule property and Lessor will permit (without any excuses what so ever) the Lessor to carry out such works within the schedule property, which are required for the general upkeep of the whole."),
            (15, "Vacant Possession",
             "The Lessor has delivered vacant possession of the Schedule Property to the Lessee this day pursuant to this deed. The Lessee shall deliver back vacant possession of the Schedule Property to the Lessor after the expiry of the period of lease fixed under this deed. The Lessee shall ensure that while vacating the schedule property all electrical fittings, plumber fittings, kitchen fittings, furniture fittings, all other fittings of the apartment as per delivered conditions and in case any damage etc found shall be replaced as per the original conditions."),
            (19, "Representation",
             "PROVIDED ALWAYS THAT whenever such an interpretation would be requisite to give fullest scope and effect legally possible, for any covenant or contract herein contained the expression 'LESSOR' shall mean and include his heirs, legal representatives, successors and assigns and the expression 'LESSEE' shall mean & include his heirs & legal representatives only."),
            (20, "Jurisdiction",
             "This Lease Deed is subject to the exclusive courts of Bangalore jurisdiction only."),
            (21, "Liability",
             "That if the Schedule Property or any part thereof be destroyed or damaged by fire (not caused by any wilful act or negligence) OR earthquake, tempest, flood, lightning, violence of any army or mob or enemies of the country or by any other irresistible force so as to render the Schedule property unfit for the purpose for which the same are let, either party shall have the option to forthwith terminate this Lease notwithstanding notice period provided in the Lease Deed. Upon this all the refundable securities shall be refunded immediately subject to any deductions / adjustments under this lease deed. The Lessor shall not be responsible or liable for any theft, loss, damage or destruction of any property of the Lessee or of any other person lying in the Schedule Property nor for any bodily injury or harm to/death of any person in the Schedule property from any cause whatsoever."),
            (22, "Possession",
             "That, the immediately at on the expiration or termination or cancellation of this agreement the Lessee shall vacate the said scheduled property and its premises without delay with all his goods and belongings. In the event of the Lessee failing and / or neglecting to remove him/herself and/or his/her articles/goods from the said & premises on expiry or sooner termination of this Agreement, the Lessor shall be entitled to recover damages at the rate of double the daily amount of compensation/rent per day and or alternatively the Lessor shall be entitled to remove the Lessee and his belongings from the Licensed premises, without recourse to the Court of Law and upon which the LESSOR shall return the remaining Security Deposit free of interest less any deduction shall be refunded one week after vacating the premises to the LESSEE."),
            (23, "Fixtures",
             "That the Lessee has seen before occupying the said scheduled property that all the sanitary and electric fittings and fixtures (details are clearly mentioned below) are in very good working condition and are satisfied that nothing is broken or missing and the Lessee on vacating the demised scheduled property premises shall restore them, in the same condition, subject to normal wear and tear."),
            (24, "MOVE-IN and MOVE-OUT Charges",
             "All costs & association charges relating to shifting in and shifting out of the community/complex shall always be borne by the LESSEE."),
            (25, "Peaceful Stay",
             "The Lessor hereby assures the Lessee that on performance of the terms and conditions contained in this deed by the Lessee, the Lessee shall quietly enjoy the Schedule Property, without any hindrance either by the Lessor or from anybody else claiming through or under the Lessor. The tenant shall not cause any disturbance and live peacefully without harming the interest of neighbors and by fulfilling the above terms and conditions regularly without default."),
            (26, "Water Charges",
             "Currently, no water charges are applicable. However, the association is in discussions regarding the introduction of water charges very soon. If any such charges are implemented in future, the same will be communicated separately and shall be payable by lessee on a monthly basis along with Rent, electricity amount."),
            (27, "Cleaning Responsibility",
             "Cleaning of the staircase area from the 1st floor up to the next (2nd) floor mid steps will be the responsibility of the respective tenant. If tenants are not agreeable to this arrangement, a cleaner will be arranged for common area maintenance (parking, staircase, common areas etc.), and the cleaning charges will be shared equally among all tenants."),
        ]:
            clause(num, title)
            for para_text in text.split("||PARA||"):
                body([(para_text.strip(), False)])
                sp()

        # Clause 14 — dynamic
        clause(14, "Termination / Expiry / Revocation of the Lease")
        body([
            (f"The agreement shall be liable to be terminated by either Lessee or the Lessor by giving minimum of {notice} of written (by mail/WhatsApp/call) notice. Since the Lessor have freshly painted the Schedule Property before handing it over to the Lessee, an amount equivalent to ONE MONTH rent (", False),
            (f"Rs.{get('monthly_rent_revised')}/-", True),
            (") prevailing at the time of the Lessee vacating the Schedule Property shall be deducted as repainting charges by the Lessor from the interest free refundable deposit of the Lessee or Lessee can also repaint the house on vacating time then there is No deduction for painting. Any damages, Electricity bills, Gas bills, water bills, Unpaid Rents and any other due charges, amount will be deducted from the LESSEE's security deposit at the time of vacating the schedule premises. If in case, Lessee not able to vacate the said premises in the given vacant date then lessor shall be entitled to receive an interest at 18% per annum on the rent for the period of extension/delay.", False),
        ])
        sp()

        # Clause 16 — dynamic
        clause(16, "Breach Of Contract")
        body([(f"In the event of the Lessee committing breach of any of the terms and conditions contained in this deed, the Lessor is at liberty to terminate the lease and seek vacant possession of the Schedule Property from the Lessee even before the expiry of the period of lease fixed under this deed by immediately giving {notice} notice. If the Lessee fails to pay the monthly rent to the Lessor, for more than two consecutive months, the Lessor reserves the right to terminate this contract and seek vacant possession of the Schedule Property from the Lessee, even before expiry of the period of lease fixed under this deed by immediately giving {notice} notice.", False)])
        sp()

        # Clause 17 — dynamic
        clause(17, "Termination Notice")
        body([(f"It is agreed that the contract entered into under this deed, can be terminated by either of the parties at any time, by giving {notice} notice in writing.", False)])
        sp()
        body([("The Rental Agreement shall be terminated under all or any of the following circumstances, namely.", False)])
        body([("   i.  By efflux of time;", False)])
        body([("   ii. In the event of breach by either party of the terms, conditions and covenants hereof;", False)])
        body([(f"   iii. By giving {notice} prior notice from either party.", False)])
        sp()

        # Clause 18 — dynamic
        clause(18, "Renewal Of Lease Deed")
        body([(f"After expiry of the period of lease, if the Lessor and the Lessee agree to renew this contract, the hike in the monthly rent payable by the Lessee to the Lessor shall be {enh_pct}% ({enh_word} percent) of the monthly rent (fixed during the first term) for these mutual discussions.", False)])
        sp()

        # SCHEDULE
        doc.add_page_break()
        heading("SCHEDULE")
        sp()
        body([
            ("The Residential Premises bearing ", False),
            (prop_full, True),
            (", Karnataka., comprising of one living room, one kitchen with Utility area, one bedroom, one bathroom and one pooja Cabinet.", False),
        ])
        sp()
        body([("IN WITNESS WHEREOF, both the Lessor and the Lessee have affixed their respective hands and signature to this agreement on this day, month and year first above written.", False)])
        sp()
        body([("WITNESSES:", True)])
        sp()
        body([(f"1) {get('owner_name')}", False)])
        body([(f"   Mob# {get('owner_mobile')}", False)])
        body([("   ", False), ("LESSOR", True)])
        sp()
        body([(f"2) {get('tenant_name')}.", False)])
        body([(f"   Aadhar ID # {get('tenant_aadhaar')}", False)])
        body([(f"   Mob # {get('tenant_mobile')}", False)])
        body([("   ", False), ("LESSEE", True)])

        # FITTINGS & FIXTURES
        doc.add_page_break()
        heading("SCHEDULE PREMISES")
        heading("FITTINGS & FIXTURE IN SCHEDULED PREMISES")
        sp()

        def fixture_row(left, right=""):
            p = doc.add_paragraph()
            run = p.add_run(f"{left:<52}{right}")
            run.font.name = "Courier New"
            run.font.size = Pt(10)

        fixture_row("Hall Area:", "Kitchen:")
        fixture_row(f"Foyer Light: {get('hall_foyer_light')}", f"Modular kitchen with Glasses: {get('kitchen_modular')}")
        fixture_row(f"Tube Light: {get('hall_tube_light')}", f"LED/Zoomer Light: {get('kitchen_led')}")
        fixture_row(f"LED/Zoomer Light: {get('hall_led_light')}", f"LED/Zoomer Light in Utility: {get('kitchen_led_utility')}")
        fixture_row(f"Curtain Rod: {get('hall_curtain_rod')}")
        fixture_row(f"Wall Mounted Bulb: {get('hall_wall_bulb')}")
        fixture_row(f"Fan: {get('hall_fan')}")
        fixture_row(f"TV Cabinet with Glasses: {get('hall_tv_cabinet')}")
        fixture_row(f"UPS Point: {get('hall_ups')}")
        fixture_row(f"Bell: {get('hall_bell')}")
        fixture_row(f"False Ceiling/POP Lights: {get('hall_pop_lights')}")
        sp()
        body([("Pooja Mandir:", True)])
        fixture_row(f"Pooja Mandir structure set-up: {get('pooja_structure')}")
        fixture_row(f"Storage Box/Drawers: {get('pooja_storage')}")
        fixture_row(f"Pooja Bulb: {get('pooja_bulb')}")
        sp()
        body([("Bedroom:", True)])
        fixture_row(f"Fan: {get('bed_fan')}", f"Tube light: {get('bed_tube')}")
        fixture_row(f"Curtain rod: {get('bed_curtain')}", f"Loft: {get('bed_loft')}")
        fixture_row(f"Sliding Lock: {get('bed_sliding_lock')}", f"Drawer Locker: {get('bed_drawer_locker')}")
        fixture_row(f"Dressing Glass: {get('bed_dressing')}", f"Bangles Rod: {get('bed_bangles_rod')}")
        fixture_row(f"Hanger Rod: {get('bed_hanger')}", f"Bulb: {get('bed_bulb')}")
        fixture_row(f"Ceiling Light: {get('bed_ceiling')}", f"Sitting Desk: {get('bed_desk')}")
        sp()
        body([("Bathroom:", True)])
        fixture_row(f"Storage Box (with 1 Mirror): {get('bath_storage')}", f"Geyser: {get('bath_geyser')}")
        fixture_row(f"Bulb: {get('bath_bulb')}", f"Ceiling Light: {get('bath_ceiling')}")
        fixture_row(f"Exhaust Fan: {get('bath_exhaust')}")
        sp()
        body([("Total Keys of Flat:", True)])
        fixture_row(f"Main Door: {get('key_main_num')} ({get('key_main_qty')})", f"Bedroom: {get('key_bedroom_num')} ({get('key_bedroom_qty')})")
        fixture_row(f"Sliding Door Lock: {get('key_sliding_num')} ({get('key_sliding_qty')})", f"Drawer Lock: {get('key_drawer_num')} ({get('key_drawer_qty')})")

        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        tenant_name_safe = get("tenant_name", "Tenant").replace(" ", "_")
        from flask import send_file
        return send_file(
            buf,
            as_attachment=True,
            download_name=f"Rental_Agreement_{tenant_name_safe}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        import traceback
        flash(f"Error generating document: {str(e)}", "danger")
        return redirect(url_for("rental_agreement_form", lead_id=lead_id))


@app.route("/lead/status", methods=["POST"])
@csrf.exempt
def update_lead_status():
    """Update lead status."""
    lead_id   = request.form.get("lead_id", "").strip()
    status    = request.form.get("status", "").strip()
    valid     = ["new", "in_discussion", "expired", "converted"]
    if not lead_id or status not in valid:
        return jsonify({"error": "Invalid"}), 400
    try:
        conn = get_db()
        conn.execute("UPDATE leads SET status=? WHERE id=?", (status, int(lead_id)))
        conn.commit()
        conn.close()
        return jsonify({"ok": True}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/lead/save", methods=["POST"])
@csrf.exempt
def save_lead():
    """Save tenant lead when they click WhatsApp contact."""
    prop_id = request.form.get("property_id", "").strip()
    name    = request.form.get("tenant_name", "").strip()
    phone   = request.form.get("tenant_phone", "").strip()

    if not prop_id or not name or not phone:
        return jsonify({"error": "Missing fields"}), 400

    try:
        conn = get_db()
        # Verify property exists
        prop = conn.execute("SELECT id FROM properties WHERE id=?", (prop_id,)).fetchone()
        if not prop:
            conn.close()
            return jsonify({"error": "Property not found"}), 404

        conn.execute(
            "INSERT INTO leads (property_id, tenant_name, tenant_phone) VALUES (?,?,?)",
            (int(prop_id), name, phone)
        )
        conn.commit()
        conn.close()
        return jsonify({"ok": True}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/leads/<int:prop_id>")
@login_required
def property_leads(prop_id):
    """Show leads for a specific property."""
    conn = get_db()
    prop = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=?",
        (prop_id, session["owner_id"])
    ).fetchone())

    if not prop:
        flash("Property not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))

    leads = conn.execute(
        "SELECT * FROM leads WHERE property_id=? ORDER BY created_at DESC",
        (prop_id,)
    ).fetchall()
    conn.close()

    return render_template("leads.html",
                           prop=prop,
                           leads=leads,
                           brand_name=session["brand_name"],
                           owner_name=session["owner_name"],
                           owner_slug=session["owner_slug"])


@app.errorhandler(404)
def page_not_found(e):
    return render_template("404.html"), 404


@app.errorhandler(429)
def rate_limit_exceeded(e):
    return render_template("login.html", rate_limited=True), 429


# ===============================================
# PUBLIC ROUTES
# ===============================================

@app.route("/")
def landing():
    return render_template("landing.html")


@app.route("/search")
def search():
    return render_template("tenant_start.html")


@app.route("/results")
def search_results():
    conn = get_db()

    prop_type = request.args.get("type", "").strip()
    beds      = request.args.get("beds", "").strip()
    min_rent  = request.args.get("min_rent", "").strip()
    max_rent  = request.args.get("max_rent", "").strip()
    group_by  = request.args.get("group", "").strip()

    query = """
        SELECT p.*,
               o.slug, o.name as owner_name, o.brand_name as owner_brand
        FROM properties p
        JOIN owners o ON p.owner_id = o.id
        WHERE o.is_active = 1
        AND (p.parent_id IS NULL OR p.parent_id = 0)
        AND (p.visibility = 'public' OR p.visibility IS NULL OR p.visibility = '')
    """
    params = []

    if prop_type:
        query += " AND p.property_type = ?"
        params.append(prop_type)

    if beds:
        if beds == "4":
            query += " AND p.beds >= 4"
        else:
            query += " AND p.beds = ?"
            params.append(int(beds))

    query += " ORDER BY p.id DESC"

    rows = conn.execute(query, params).fetchall()
    conn.close()

    conn2 = get_db()
    properties = []
    for r in rows:
        p = parse_property(r)
        p["display_title"] = p["title"]
        # For complexes, load sub-unit summary
        if p.get("is_complex"):
            units = conn2.execute(
                "SELECT title, beds, status FROM properties WHERE parent_id=? ORDER BY id",
                (p["id"],)
            ).fetchall()
            p["unit_count"] = len(units)
            p["unit_available"] = sum(1 for u in units if u["status"] == "Available")
            # Show unit names (up to 5 then ...)
            names = [u["title"] for u in units if u["title"]]
            if len(names) <= 5:
                p["unit_labels"] = ", ".join(names)
            else:
                p["unit_labels"] = ", ".join(names[:5]) + "..."
        else:
            if min_rent or max_rent:
                rent_num = int("".join(filter(str.isdigit, p["rent"] or "0")))
                if min_rent and rent_num < int(min_rent):
                    continue
                if max_rent and rent_num > int(max_rent):
                    continue
        properties.append(p)
    conn2.close()

    grouped = {}
    if group_by == "type":
        for p in properties:
            key = p.get("property_type") or "Other"
            grouped.setdefault(key, []).append(p)

    elif group_by == "owner":
        for p in properties:
            key = p.get("owner_brand") or p.get("owner_name") or "Unknown"
            grouped.setdefault(key, []).append(p)

    elif group_by == "beds":
        order = ["Studio", "1 Bedroom", "2 Bedrooms", "3 Bedrooms", "4+ Bedrooms"]
        for p in properties:
            b = p.get("beds") or 0
            if b == 0:   key = "Studio"
            elif b == 1: key = "1 Bedroom"
            elif b == 2: key = "2 Bedrooms"
            elif b == 3: key = "3 Bedrooms"
            else:        key = "4+ Bedrooms"
            grouped.setdefault(key, []).append(p)
        grouped = {k: grouped[k] for k in order if k in grouped}

    elif group_by == "status":
        order = ["Available", "Rented", "Maintenance"]
        for p in properties:
            key = p.get("status") or "Other"
            grouped.setdefault(key, []).append(p)
        grouped = {k: grouped[k] for k in order if k in grouped}

    return render_template("search_results.html",
                           properties=properties,
                           grouped=grouped,
                           group_by=group_by,
                           prop_type=prop_type,
                           beds=beds,
                           min_rent=min_rent,
                           max_rent=max_rent)


RESERVED = {"admin", "login", "logout", "dashboard", "create", "edit",
            "delete", "enquiries", "enquire", "static", "favicon.ico",
            "search", "results"}


@app.route("/<slug>")
def owner_portal(slug):
    if slug in RESERVED:
        return redirect(url_for("landing"))

    conn = get_db()
    owner = conn.execute(
        "SELECT * FROM owners WHERE slug=? AND is_active=1", (slug,)
    ).fetchone()

    if not owner:
        conn.close()
        return render_template("404.html"), 404

    properties = [parse_property(r) for r in conn.execute("""
        SELECT * FROM properties
        WHERE owner_id=? AND (parent_id IS NULL OR parent_id=0)
        ORDER BY id DESC
    """, (owner["id"],)).fetchall()]
    conn.close()

    return render_template("owner_portal.html", owner=dict(owner), properties=properties)


@app.route("/<slug>/property/<int:prop_id>")
def property_detail(slug, prop_id):
    conn = get_db()
    owner = conn.execute(
        "SELECT * FROM owners WHERE slug=? AND is_active=1", (slug,)
    ).fetchone()
    if not owner:
        conn.close()
        return render_template("404.html"), 404

    prop = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=?",
        (prop_id, owner["id"])
    ).fetchone())

    if not prop:
        conn.close()
        return render_template("404.html"), 404

    # Hidden properties return 404 even on direct URL
    if prop.get("visibility") == "hidden":
        conn.close()
        return render_template("404.html"), 404

    sub_units = []
    if prop.get("is_complex"):
        sub_units = [parse_property(r) for r in conn.execute(
            "SELECT * FROM properties WHERE parent_id=? ORDER BY id",
            (prop_id,)
        ).fetchall()]

    conn.close()
    return render_template("property_detail.html",
                           prop=prop, owner=dict(owner),
                           sub_units=sub_units, slug=slug,
                           whatsapp_phone=clean_phone(owner["phone"] or ""))


# ===============================================
# OWNER AUTH
# ===============================================

@app.route("/login", methods=["GET", "POST"])
@limiter.limit("10 per minute")
def login():
    if "owner_id" in session:
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        password = request.form.get("password", "")
        conn     = get_db()
        owner    = conn.execute(
            "SELECT * FROM owners WHERE LOWER(username)=? AND is_active=1", (username,)
        ).fetchone()
        conn.close()

        if owner and check_password_hash(owner["password"], password):
            session.clear()
            session["owner_id"]   = owner["id"]
            session["owner_name"] = owner["name"]
            session["owner_slug"] = owner["slug"]
            session["brand_name"] = owner["brand_name"]
            flash(f"Welcome back, {owner['name']}!", "success")
            return redirect(url_for("dashboard"))
        else:
            flash("Incorrect username or password.", "danger")

    return render_template("login.html")


@app.route("/logout", methods=["GET","POST"])
def logout():
    session.clear()
    resp = make_response(redirect("/"))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"]        = "no-cache"
    resp.headers["Expires"]       = "0"
    return resp


# ===============================================
# OWNER DASHBOARD
# ===============================================

@app.route("/dashboard")
@login_required
def dashboard():
    conn = get_db()

    owner = conn.execute(
        "SELECT * FROM owners WHERE id=? AND is_active=1",
        (session["owner_id"],)
    ).fetchone()

    if not owner:
        conn.close()
        session.clear()
        flash("Your account is no longer active.", "danger")
        return redirect(url_for("login"))

    session["brand_name"] = owner["brand_name"]
    session["owner_slug"] = owner["slug"]
    session["owner_name"] = owner["name"]

    all_props = [parse_property(r) for r in conn.execute(
        "SELECT * FROM properties WHERE owner_id=? ORDER BY id DESC",
        (session["owner_id"],)
    ).fetchall()]

    # Load lead counts per property
    lead_counts = {}
    for row in conn.execute(
        "SELECT property_id, COUNT(*) as cnt FROM leads WHERE property_id IN "
        "(SELECT id FROM properties WHERE owner_id=?) GROUP BY property_id",
        (session["owner_id"],)
    ).fetchall():
        lead_counts[row["property_id"]] = row["cnt"]

    # For complexes, attach unit summary
    for p in all_props:
        if p.get("is_complex"):
            units = conn.execute(
                "SELECT title, beds, status FROM properties WHERE parent_id=? ORDER BY id",
                (p["id"],)
            ).fetchall()
            p["unit_count"] = len(units)
            p["unit_available"] = sum(1 for u in units if u["status"] == "Available")
            names = [u["title"] for u in units if u["title"]]
            p["unit_names"] = " · ".join(names[:5]) + ("..." if len(names) > 5 else "")
        else:
            p["unit_count"] = 0

    properties = all_props

    enquiries = {}
    for row in conn.execute(
        "SELECT property_id, COUNT(*) as cnt FROM enquiries WHERE property_id IN "
        "(SELECT id FROM properties WHERE owner_id=?) GROUP BY property_id",
        (session["owner_id"],)
    ).fetchall():
        enquiries[row["property_id"]] = row["cnt"]

    conn.close()

    resp = make_response(render_template("dashboard.html",
                           properties=properties,
                           enquiries=enquiries,
                           lead_counts=lead_counts,
                           owner_name=session["owner_name"],
                           brand_name=session["brand_name"],
                           owner_slug=session["owner_slug"]))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"]        = "no-cache"
    resp.headers["Expires"]       = "0"
    return resp



@app.route("/create/rent", methods=["GET", "POST"])
@login_required
def create_rent():
    return create_property_handler("rent")

@app.route("/create/sale", methods=["GET", "POST"])
@login_required
def create_sale():
    return create_property_handler("sale")

def create_property_handler(listing_type):
    if request.method == "POST":
        is_complex     = request.form.get("is_complex") == "1"
        gallery_raw    = request.form.get("gallery_images", "")
        gallery_images = json.dumps([g.strip() for g in gallery_raw.splitlines() if g.strip()])
        images_list    = [g.strip() for g in gallery_raw.splitlines() if g.strip()]
        first_image    = images_list[0] if images_list else ""
        visibility     = request.form.get("visibility", "public")
        sale_price     = request.form.get("sale_price", "")
        contact_for_price = 1 if request.form.get("contact_for_price") else 0

        conn = get_db()

        if is_complex:
            amenities = json.dumps([])
            c_listing_type = listing_type
            c_visibility   = visibility
            conn.execute("""INSERT INTO properties
                (owner_id, title, location, rent, status, beds, baths, sqft, image,
                 property_type, security_deposit, facing, maintenance, floor, built,
                 description, amenities, gallery_images, is_complex,
                 visibility, listing_type, sale_price, contact_for_price)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,1,?,?,?,?)""", (
                session["owner_id"],
                request.form.get("title"), request.form.get("location"),
                "", "Available", 0, 0, 0, first_image,
                request.form.get("property_type"),
                "", request.form.get("facing"), "", "",
                request.form.get("built"), request.form.get("description"),
                amenities, gallery_images,
                c_visibility, c_listing_type, sale_price, contact_for_price
            ))
            complex_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]

            unit_names        = request.form.getlist("unit_title[]")
            unit_rents        = request.form.getlist("unit_rent[]")
            unit_sale_prices  = request.form.getlist("unit_sale_price[]")
            unit_cfp          = request.form.getlist("unit_contact_for_price[]")
            unit_beds         = request.form.getlist("unit_beds[]")
            unit_baths        = request.form.getlist("unit_baths[]")
            unit_sqfts        = request.form.getlist("unit_sqft[]")
            unit_statuses     = request.form.getlist("unit_status[]")
            unit_floors       = request.form.getlist("unit_floor[]")
            unit_galleries    = request.form.getlist("unit_gallery_images[]")
            unit_sec_deposits = request.form.getlist("unit_security_deposit[]")
            unit_maintenances = request.form.getlist("unit_maintenance[]")
            unit_visibilities = request.form.getlist("unit_visibility[]")

            for i, name in enumerate(unit_names):
                if not name.strip(): continue
                ugallery_raw  = unit_galleries[i] if i < len(unit_galleries) else ""
                ugallery_list = [g.strip() for g in ugallery_raw.splitlines() if g.strip()]
                ugallery_json = json.dumps(ugallery_list)
                ufirst_image  = ugallery_list[0] if ugallery_list else first_image
                u_sale_price  = unit_sale_prices[i] if i < len(unit_sale_prices) else ""
                u_cfp         = 1 if (i < len(unit_cfp) and unit_cfp[i]) else 0
                u_vis         = unit_visibilities[i] if i < len(unit_visibilities) else "public"
                conn.execute("""INSERT INTO properties
                    (owner_id, parent_id, title, location, rent, status, beds, baths,
                     sqft, image, property_type, floor, security_deposit, maintenance,
                     description, amenities, gallery_images, is_complex,
                     listing_type, sale_price, contact_for_price, visibility)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,0,?,?,?,?)""", (
                    session["owner_id"], complex_id, name.strip(),
                    request.form.get("location"),
                    unit_rents[i] if i < len(unit_rents) else "",
                    unit_statuses[i] if i < len(unit_statuses) else "Available",
                    int(unit_beds[i]) if i < len(unit_beds) and unit_beds[i] else 0,
                    int(unit_baths[i]) if i < len(unit_baths) and unit_baths[i] else 0,
                    int(unit_sqfts[i]) if i < len(unit_sqfts) and unit_sqfts[i] else 0,
                    ufirst_image, request.form.get("property_type"),
                    unit_floors[i] if i < len(unit_floors) else "",
                    unit_sec_deposits[i] if i < len(unit_sec_deposits) else "",
                    unit_maintenances[i] if i < len(unit_maintenances) else "",
                    "", json.dumps([]), ugallery_json,
                    listing_type, u_sale_price, u_cfp, u_vis
                ))
            conn.commit()
            conn.close()
            flash(f"Complex created with {len([n for n in unit_names if n.strip()])} units!", "success")
            return redirect(url_for("dashboard"))

        else:
            amenities_raw  = request.form.get("amenities", "")
            amenities      = json.dumps([a.strip() for a in amenities_raw.split(",") if a.strip()])
            conn.execute("""INSERT INTO properties
                (owner_id, title, location, rent, status, beds, baths, sqft, image,
                 property_type, security_deposit, facing, maintenance, floor, built,
                 description, amenities, gallery_images, is_complex,
                 listing_type, sale_price, contact_for_price, visibility)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,0,?,?,?,?)""", (
                session["owner_id"],
                request.form.get("title"), request.form.get("location"),
                request.form.get("rent") if listing_type == "rent" else "",
                request.form.get("status"),
                request.form.get("beds") or 0,
                request.form.get("baths") or 0,
                request.form.get("sqft") or 0,
                first_image,
                request.form.get("property_type"),
                request.form.get("security_deposit"),
                request.form.get("facing"),
                request.form.get("maintenance"),
                request.form.get("floor"),
                request.form.get("built"),
                request.form.get("description"),
                amenities, gallery_images,
                listing_type, sale_price, contact_for_price, visibility
            ))
            conn.commit()
            conn.close()
            flash(f"'{request.form.get('title')}' listed successfully!", "success")
            return redirect(url_for("dashboard"))

    return render_template("create_property.html",
                           brand_name=session["brand_name"],
                           listing_type=listing_type)

@app.route("/create", methods=["GET"])
@login_required
def create_property():
    return render_template("create_landing.html", brand_name=session["brand_name"])


@app.route("/create_post", methods=["POST"])
@login_required
def create_property_post():
    if request.method == "POST":
        is_complex     = request.form.get("is_complex") == "1"
        gallery_raw    = request.form.get("gallery_images", "")
        gallery_images = json.dumps([g.strip() for g in gallery_raw.splitlines() if g.strip()])
        images_list    = [g.strip() for g in gallery_raw.splitlines() if g.strip()]
        first_image    = images_list[0] if images_list else ""

        conn = get_db()

        if is_complex:
            # Complex: minimal fields, no rent/beds/baths
            amenities = json.dumps([])
            c_visibility = request.form.get("visibility", "public")
            c_listing_type = request.form.get("listing_type", "rent")
            conn.execute("""INSERT INTO properties
                (owner_id, title, location, rent, status, beds, baths, sqft, image,
                 property_type, security_deposit, facing, maintenance, floor, built,
                 description, amenities, gallery_images, is_complex,
                 visibility, listing_type)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,1,?,?)""", (
                session["owner_id"],
                request.form.get("title"),
                request.form.get("location"),
                "", "Available", 0, 0, 0, first_image,
                request.form.get("property_type"),
                "", request.form.get("facing"), "", "",
                request.form.get("built"),
                request.form.get("description"),
                amenities, gallery_images,
                c_visibility, c_listing_type
            ))
            complex_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]

            # Save inline units submitted with complex form
            unit_names    = request.form.getlist("unit_title[]")
            unit_rents    = request.form.getlist("unit_rent[]")
            unit_beds     = request.form.getlist("unit_beds[]")
            unit_baths    = request.form.getlist("unit_baths[]")
            unit_sqfts    = request.form.getlist("unit_sqft[]")
            unit_statuses = request.form.getlist("unit_status[]")
            unit_floors   = request.form.getlist("unit_floor[]")
            unit_galleries       = request.form.getlist("unit_gallery_images[]")
            unit_sec_deposits    = request.form.getlist("unit_security_deposit[]")
            unit_maintenances    = request.form.getlist("unit_maintenance[]")

            for i, name in enumerate(unit_names):
                if not name.strip():
                    continue
                # Parse unit gallery images
                ugallery_raw = unit_galleries[i] if i < len(unit_galleries) else ""
                ugallery_list = [g.strip() for g in ugallery_raw.splitlines() if g.strip()]
                ugallery_json = json.dumps(ugallery_list)
                ufirst_image = ugallery_list[0] if ugallery_list else first_image

                u_listing_type      = c_listing_type
                u_sale_prices       = request.form.getlist("unit_sale_price[]")
                u_contact_for_price = request.form.getlist("unit_contact_for_price[]")
                u_visibilities      = request.form.getlist("unit_visibility[]")
                u_sale_price        = u_sale_prices[i] if i < len(u_sale_prices) else ""
                u_cfp               = 1 if (i < len(u_contact_for_price) and u_contact_for_price[i]) else 0
                u_vis               = u_visibilities[i] if i < len(u_visibilities) else "public"
                conn.execute("""INSERT INTO properties
                    (owner_id, parent_id, title, location, rent, status, beds, baths,
                     sqft, image, property_type, floor, security_deposit, maintenance,
                     description, amenities, gallery_images, is_complex,
                     listing_type, sale_price, contact_for_price, visibility)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,0,?,?,?,?)""", (
                    session["owner_id"], complex_id,
                    name.strip(),
                    request.form.get("location"),
                    unit_rents[i] if i < len(unit_rents) else "",
                    unit_statuses[i] if i < len(unit_statuses) else "Available",
                    int(unit_beds[i]) if i < len(unit_beds) and unit_beds[i] else 0,
                    int(unit_baths[i]) if i < len(unit_baths) and unit_baths[i] else 0,
                    int(unit_sqfts[i]) if i < len(unit_sqfts) and unit_sqfts[i] else 0,
                    ufirst_image, request.form.get("property_type"),
                    unit_floors[i] if i < len(unit_floors) else "",
                    unit_sec_deposits[i] if i < len(unit_sec_deposits) else "",
                    unit_maintenances[i] if i < len(unit_maintenances) else "",
                    "", json.dumps([]), ugallery_json,
                    u_listing_type, u_sale_price, u_cfp, u_vis
                ))

            conn.commit()
            conn.close()
            flash(f"Complex '{request.form.get('title')}' created with {len([n for n in unit_names if n.strip()])} units!", "success")
            return redirect(url_for("dashboard"))

        else:
            # Single property — all fields
            amenities_raw  = request.form.get("amenities", "")
            amenities      = json.dumps([a.strip() for a in amenities_raw.split(",") if a.strip()])
            listing_type     = request.form.get("listing_type", "rent")
            sale_price       = request.form.get("sale_price", "")
            contact_for_price= 1 if request.form.get("contact_for_price") else 0
            visibility       = request.form.get("visibility", "public")
            conn.execute("""INSERT INTO properties
                (owner_id, title, location, rent, status, beds, baths, sqft, image,
                 property_type, security_deposit, facing, maintenance, floor, built,
                 description, amenities, gallery_images, is_complex,
                 listing_type, sale_price, contact_for_price, visibility)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,0,?,?,?,?)""", (
                session["owner_id"],
                request.form.get("title"),
                request.form.get("location"),
                request.form.get("rent") if listing_type == "rent" else "",
                request.form.get("status"),
                request.form.get("beds") or 0,
                request.form.get("baths") or 0,
                request.form.get("sqft") or 0,
                first_image,
                request.form.get("property_type"),
                request.form.get("security_deposit"),
                request.form.get("facing"),
                request.form.get("maintenance"),
                request.form.get("floor"),
                request.form.get("built"),
                request.form.get("description"),
                amenities, gallery_images,
                listing_type, sale_price, contact_for_price, visibility
            ))
            conn.commit()
            conn.close()
            flash(f"'{request.form.get('title')}' created successfully!", "success")
            return redirect(url_for("dashboard"))

    return render_template("create_property.html",
                           brand_name=session["brand_name"],
                           listing_type="rent")



@app.route("/add-unit/<int:complex_id>", methods=["GET", "POST"])
@login_required
def add_unit(complex_id):
    """Add a unit to an existing complex property."""
    conn = get_db()
    parent = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=? AND is_complex=1",
        (complex_id, session["owner_id"])
    ).fetchone())

    if not parent:
        flash("Complex not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        gallery_raw    = request.form.get("gallery_images", "")
        amenities_raw  = request.form.get("amenities", "")
        gallery_images = json.dumps([g.strip() for g in gallery_raw.splitlines() if g.strip()])
        amenities      = json.dumps([a.strip() for a in amenities_raw.split(",") if a.strip()])
        images_list    = [g.strip() for g in gallery_raw.splitlines() if g.strip()]
        first_image    = images_list[0] if images_list else parent.get("image", "")

        conn.execute("""INSERT INTO properties
            (owner_id, parent_id, title, location, rent, status, beds, baths, sqft,
             image, property_type, security_deposit, facing, maintenance, floor,
             description, amenities, gallery_images, is_complex)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,0)""", (
            session["owner_id"],
            complex_id,
            request.form.get("title"),
            parent["location"],   # inherit parent location
            request.form.get("rent"),
            request.form.get("status", "Available"),
            request.form.get("beds") or 0,
            request.form.get("baths") or 0,
            request.form.get("sqft") or 0,
            first_image,
            request.form.get("property_type") or parent.get("property_type"),
            request.form.get("security_deposit"),
            request.form.get("facing"),
            request.form.get("maintenance"),
            request.form.get("floor"),
            request.form.get("description"),
            amenities,
            gallery_images
        ))
        conn.commit()
        conn.close()
        flash(f"Unit added to {parent['title']}!", "success")
        return redirect(url_for("dashboard"))

    conn.close()
    return render_template("add_unit.html",
                           parent=parent,
                           brand_name=session["brand_name"])



@app.route("/edit-complex/<int:prop_id>", methods=["GET", "POST"])
@login_required
def edit_complex(prop_id):
    conn = get_db()
    prop = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=? AND is_complex=1",
        (prop_id, session["owner_id"])
    ).fetchone())

    if not prop:
        flash("Complex not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        gallery_raw    = request.form.get("gallery_images", "")
        gallery_images = json.dumps([g.strip() for g in gallery_raw.splitlines() if g.strip()])
        images_list    = [g.strip() for g in gallery_raw.splitlines() if g.strip()]
        first_image    = images_list[0] if images_list else prop.get("image", "")

        ec_visibility   = request.form.get("visibility", "public")
        ec_listing_type = request.form.get("listing_type", "rent")
        conn.execute("""UPDATE properties SET
            title=?, location=?, property_type=?, facing=?, built=?,
            description=?, image=?, gallery_images=?,
            visibility=?, listing_type=?
            WHERE id=? AND owner_id=?""", (
            request.form.get("title"),
            request.form.get("location"),
            request.form.get("property_type"),
            request.form.get("facing"),
            request.form.get("built"),
            request.form.get("description"),
            first_image, gallery_images,
            ec_visibility, ec_listing_type,
            prop_id, session["owner_id"]
        ))
        conn.commit()
        conn.close()
        flash(f"'{request.form.get('title')}' updated!", "success")
        return redirect(url_for("dashboard"))

    sub_units = [parse_property(r) for r in conn.execute(
        "SELECT * FROM properties WHERE parent_id=? ORDER BY id",
        (prop_id,)
    ).fetchall()]
    conn.close()
    prop["gallery_images_str"] = "\n".join(prop["gallery_images"])
    return render_template("edit_complex.html", prop=prop, sub_units=sub_units,
                           brand_name=session["brand_name"])


@app.route("/edit/<int:prop_id>", methods=["GET", "POST"])
@login_required
def edit_property(prop_id):
    conn = get_db()
    prop = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=?",
        (prop_id, session["owner_id"])
    ).fetchone())

    if not prop:
        flash("Property not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))

    # Units go to edit_unit, complexes go to edit_complex
    if prop.get("parent_id"):
        conn.close()
        return redirect(url_for("edit_unit", unit_id=prop_id))
    if prop.get("is_complex"):
        conn.close()
        return redirect(url_for("edit_complex", prop_id=prop_id))

    if request.method == "POST":
        amenities_raw  = request.form.get("amenities", "")
        gallery_raw    = request.form.get("gallery_images", "")
        amenities      = json.dumps([a.strip() for a in amenities_raw.split(",") if a.strip()])
        gallery_images = json.dumps([g.strip() for g in gallery_raw.splitlines() if g.strip()])
        images_list    = [g.strip() for g in gallery_raw.splitlines() if g.strip()]
        first_image    = images_list[0] if images_list else prop.get("image", "")

        e_listing_type      = request.form.get("listing_type", "rent")
        e_sale_price        = request.form.get("sale_price", "")
        e_contact_for_price = 1 if request.form.get("contact_for_price") else 0
        e_visibility        = request.form.get("visibility", "public")
        conn.execute("""UPDATE properties SET
            title=?, location=?, rent=?, status=?, beds=?, baths=?, sqft=?, image=?,
            floor=?, built=?, property_type=?, security_deposit=?, facing=?,
            maintenance=?, description=?, amenities=?, gallery_images=?,
            listing_type=?, sale_price=?, contact_for_price=?, visibility=?
            WHERE id=? AND owner_id=?""", (
            request.form.get("title"),
            request.form.get("location"),
            request.form.get("rent") if e_listing_type == "rent" else "",
            request.form.get("status"),
            request.form.get("beds") or 0,
            request.form.get("baths") or 0,
            request.form.get("sqft") or 0,
            first_image,
            request.form.get("floor"),
            request.form.get("built"),
            request.form.get("property_type"),
            request.form.get("security_deposit"),
            request.form.get("facing"),
            request.form.get("maintenance"),
            request.form.get("description"),
            amenities, gallery_images,
            e_listing_type, e_sale_price, e_contact_for_price, e_visibility,
            prop_id, session["owner_id"]
        ))
        conn.commit()
        conn.close()
        flash(f"'{request.form.get('title')}' updated successfully!", "success")
        return redirect(url_for("dashboard"))

    sub_units = []
    if prop.get("is_complex"):
        sub_units = [parse_property(r) for r in conn.execute(
            "SELECT * FROM properties WHERE parent_id=? ORDER BY id",
            (prop_id,)
        ).fetchall()]
        for u in sub_units:
            u["gallery_images_str"] = "\n".join(u["gallery_images"])

    conn.close()
    prop["amenities_str"]      = ", ".join(prop["amenities"])
    prop["gallery_images_str"] = "\n".join(prop["gallery_images"])
    return render_template("edit_property.html", prop=prop, sub_units=sub_units,
                           brand_name=session["brand_name"])



@app.route("/edit-unit/<int:unit_id>", methods=["GET", "POST"])
@login_required
def edit_unit(unit_id):
    conn = get_db()
    unit = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=? AND parent_id IS NOT NULL AND parent_id != 0",
        (unit_id, session["owner_id"])
    ).fetchone())

    if not unit:
        flash("Unit not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))

    parent = parse_property(conn.execute(
        "SELECT * FROM properties WHERE id=?", (unit["parent_id"],)
    ).fetchone())

    if request.method == "POST":
        gallery_raw    = request.form.get("gallery_images", "")
        amenities_raw  = request.form.get("amenities", "")
        gallery_images = json.dumps([g.strip() for g in gallery_raw.splitlines() if g.strip()])
        amenities      = json.dumps([a.strip() for a in amenities_raw.split(",") if a.strip()])
        images_list    = [g.strip() for g in gallery_raw.splitlines() if g.strip()]
        first_image    = images_list[0] if images_list else unit.get("image", "")

        eu_listing_type      = request.form.get("listing_type", "rent")
        eu_sale_price        = request.form.get("sale_price", "")
        eu_contact_for_price = 1 if request.form.get("contact_for_price") else 0
        eu_visibility        = request.form.get("visibility", "public")
        conn.execute("""UPDATE properties SET
            title=?, rent=?, status=?, beds=?, baths=?, sqft=?, image=?,
            floor=?, facing=?, security_deposit=?, maintenance=?,
            description=?, amenities=?, gallery_images=?,
            listing_type=?, sale_price=?, contact_for_price=?, visibility=?
            WHERE id=? AND owner_id=?""", (
            request.form.get("title"),
            request.form.get("rent") if eu_listing_type == "rent" else "",
            request.form.get("status"),
            request.form.get("beds") or 0,
            request.form.get("baths") or 0,
            request.form.get("sqft") or 0,
            first_image,
            request.form.get("floor"),
            request.form.get("facing"),
            request.form.get("security_deposit"),
            request.form.get("maintenance"),
            request.form.get("description"),
            amenities, gallery_images,
            eu_listing_type, eu_sale_price, eu_contact_for_price, eu_visibility,
            unit_id, session["owner_id"]
        ))
        conn.commit()
        conn.close()
        flash(f"Unit updated!", "success")
        from_url = request.args.get("from", "")
        if from_url == "/dashboard":
            return redirect(url_for("dashboard"))
        return redirect(url_for("edit_complex", prop_id=unit["parent_id"]))

    conn.close()
    unit["gallery_images_str"] = "\n".join(unit["gallery_images"])
    unit["amenities_str"] = ", ".join(unit["amenities"])
    return render_template("edit_unit.html", unit=unit, parent=parent,
                           brand_name=session["brand_name"])


@app.route("/delete/<int:prop_id>", methods=["POST"])
@login_required
def delete_property(prop_id):
    conn = get_db()
    prop = conn.execute(
        "SELECT * FROM properties WHERE id=? AND owner_id=?",
        (prop_id, session["owner_id"])
    ).fetchone()
    if not prop:
        flash("Property not found.", "danger")
        conn.close()
        return redirect(url_for("dashboard"))
    title = prop["title"]
    sub_ids = [r["id"] for r in conn.execute(
        "SELECT id FROM properties WHERE parent_id=?", (prop_id,)
    ).fetchall()]
    for sid in sub_ids:
        conn.execute("DELETE FROM enquiries WHERE property_id=?", (sid,))
        conn.execute("DELETE FROM properties WHERE id=?", (sid,))
    conn.execute("DELETE FROM enquiries WHERE property_id=?", (prop_id,))
    conn.execute("DELETE FROM properties WHERE id=? AND owner_id=?",
                 (prop_id, session["owner_id"]))
    conn.commit()
    conn.close()
    flash(f"'{title}' deleted.", "info")
    return redirect(url_for("dashboard"))


@app.route("/enquiries")
@login_required
def enquiries():
    conn = get_db()
    rows = conn.execute("""
        SELECT e.*, p.title as property_title
        FROM enquiries e JOIN properties p ON e.property_id = p.id
        WHERE p.owner_id=? ORDER BY e.created_at DESC
    """, (session["owner_id"],)).fetchall()
    conn.close()
    return render_template("enquiries.html", enquiries=rows,
                           owner_name=session["owner_name"],
                           brand_name=session["brand_name"])


# ===============================================
# ADMIN AUTH
# ===============================================

@app.route("/admin/login", methods=["GET", "POST"])
@limiter.limit("5 per minute")
def admin_login():
    if "admin_id" in session:
        return redirect(url_for("admin_dashboard"))

    if request.method == "POST":
        email    = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        conn     = get_db()
        admin    = conn.execute(
            "SELECT * FROM admins WHERE LOWER(email)=?", (email,)
        ).fetchone()
        conn.close()

        if admin and check_password_hash(admin["password"], password):
            session.clear()
            session["admin_id"]   = admin["id"]
            session["admin_name"] = admin["name"]
            flash(f"Welcome, {admin['name']}!", "success")
            return redirect(url_for("admin_dashboard"))
        else:
            flash("Incorrect admin credentials.", "danger")

    return render_template("admin_login.html")


@app.route("/admin/logout", methods=["GET","POST"])
def admin_logout():
    session.clear()
    resp = make_response(redirect("/"))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"]        = "no-cache"
    resp.headers["Expires"]       = "0"
    return resp


# ===============================================
# ADMIN DASHBOARD
# ===============================================

@app.route("/admin")
@app.route("/admin/dashboard")
@admin_required
def admin_dashboard():
    conn = get_db()
    owners           = conn.execute("SELECT * FROM owners ORDER BY created_at DESC").fetchall()
    total_owners     = conn.execute("SELECT COUNT(*) FROM owners").fetchone()[0]
    active_owners    = conn.execute("SELECT COUNT(*) FROM owners WHERE is_active=1").fetchone()[0]
    total_properties = conn.execute("SELECT COUNT(*) FROM properties").fetchone()[0]
    total_enquiries  = conn.execute("SELECT COUNT(*) FROM enquiries").fetchone()[0]

    prop_counts = {}
    for row in conn.execute(
        "SELECT owner_id, COUNT(*) as cnt FROM properties GROUP BY owner_id"
    ).fetchall():
        prop_counts[row["owner_id"]] = row["cnt"]

    enq_counts = {}
    for row in conn.execute("""
        SELECT p.owner_id, COUNT(*) as cnt FROM enquiries e
        JOIN properties p ON e.property_id = p.id GROUP BY p.owner_id
    """).fetchall():
        enq_counts[row["owner_id"]] = row["cnt"]

    conn.close()

    resp = make_response(render_template("admin_dashboard.html",
                           owners=owners,
                           prop_counts=prop_counts,
                           enq_counts=enq_counts,
                           total_owners=total_owners,
                           active_owners=active_owners,
                           total_properties=total_properties,
                           total_enquiries=total_enquiries,
                           admin_name=session["admin_name"]))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"]        = "no-cache"
    resp.headers["Expires"]       = "0"
    return resp


@app.route("/admin/create-owner", methods=["GET", "POST"])
@admin_required
@csrf.exempt
def admin_create_owner():
    if request.method == "POST":
        name       = request.form.get("name", "").strip()
        username   = request.form.get("username", "").strip().lower()
        email      = request.form.get("email", "").strip().lower()
        password   = request.form.get("password", "").strip()
        phone_raw  = request.form.get("phone", "").strip()
        brand_name = request.form.get("brand_name", "").strip()
        slug       = request.form.get("slug", "").strip().lower().replace(" ", "-")

        if not all([name, username, email, password, brand_name, slug]):
            flash("All fields are required.", "danger")
            return render_template("admin_create_owner.html",
                                   admin_name=session["admin_name"])

        # Validate email
        email_ok, email_err = validate_email(email)
        if not email_ok:
            flash(email_err, "danger")
            return render_template("admin_create_owner.html",
                                   admin_name=session["admin_name"])

        # Validate phone if provided
        if phone_raw:
            phone_ok, phone_err = validate_phone(phone_raw)
            if not phone_ok:
                flash(phone_err, "danger")
                return render_template("admin_create_owner.html",
                                       admin_name=session["admin_name"])

        phone = clean_phone(phone_raw)

        conn = get_db()
        existing = conn.execute(
            "SELECT id FROM owners WHERE LOWER(username)=? OR LOWER(email)=? OR slug=?",
            (username, email, slug)
        ).fetchone()

        if existing:
            flash("An owner with that username, email or URL already exists.", "danger")
            conn.close()
            return render_template("admin_create_owner.html",
                                   admin_name=session["admin_name"])

        conn.execute("""INSERT INTO owners
            (name, username, email, password, phone, slug, brand_name, is_active)
            VALUES (?,?,?,?,?,?,?,1)""",
            (name, username, email, generate_password_hash(password), phone, slug, brand_name))
        conn.commit()
        conn.close()
        flash(f"Owner '{name}' created successfully!", "success")

        # Send welcome email with credentials
        email_sent = send_welcome_email(name, email, username, password, slug)
        if email_sent:
            flash(f"Welcome email sent to {email}.", "info")
        else:
            flash("Account created but welcome email could not be sent. Share credentials manually.", "warning")

        return redirect(url_for("admin_dashboard"))

    return render_template("admin_create_owner.html", admin_name=session["admin_name"])


@app.route("/admin/owner/<int:owner_id>")
@admin_required
@csrf.exempt
def admin_view_owner(owner_id):
    conn = get_db()
    owner = conn.execute("SELECT * FROM owners WHERE id=?", (owner_id,)).fetchone()
    if not owner:
        flash("Owner not found.", "danger")
        conn.close()
        return redirect(url_for("admin_dashboard"))

    properties = [parse_property(r) for r in conn.execute(
        "SELECT * FROM properties WHERE owner_id=? ORDER BY id DESC", (owner_id,)
    ).fetchall()]

    enqs = conn.execute("""
        SELECT e.*, p.title as property_title FROM enquiries e
        JOIN properties p ON e.property_id = p.id
        WHERE p.owner_id=? ORDER BY e.created_at DESC
    """, (owner_id,)).fetchall()

    conn.close()
    return render_template("admin_view_owner.html",
                           owner=dict(owner),
                           properties=properties,
                           enquiries=enqs,
                           admin_name=session["admin_name"])


@app.route("/admin/edit-owner/<int:owner_id>", methods=["GET", "POST"])
@admin_required
@csrf.exempt
def admin_edit_owner(owner_id):
    conn = get_db()
    owner = conn.execute("SELECT * FROM owners WHERE id=?", (owner_id,)).fetchone()
    if not owner:
        flash("Owner not found.", "danger")
        conn.close()
        return redirect(url_for("admin_dashboard"))

    if request.method == "POST":
        name       = request.form.get("name", "").strip()
        email      = request.form.get("email", "").strip().lower()
        phone_raw  = request.form.get("phone", "").strip()
        brand_name = request.form.get("brand_name", "").strip()
        slug       = request.form.get("slug", "").strip().lower().replace(" ", "-")

        if not all([name, email, brand_name, slug]):
            flash("All fields are required.", "danger")
            conn.close()
            return redirect(url_for("admin_edit_owner", owner_id=owner_id))

        # Validate email
        email_ok, email_err = validate_email(email)
        if not email_ok:
            flash(email_err, "danger")
            conn.close()
            return redirect(url_for("admin_edit_owner", owner_id=owner_id))

        # Validate phone if provided
        if phone_raw:
            phone_ok, phone_err = validate_phone(phone_raw)
            if not phone_ok:
                flash(phone_err, "danger")
                conn.close()
                return redirect(url_for("admin_edit_owner", owner_id=owner_id))

        phone = clean_phone(phone_raw)

        conflict = conn.execute(
            "SELECT id FROM owners WHERE (LOWER(email)=? OR slug=?) AND id != ?",
            (email, slug, owner_id)
        ).fetchone()

        if conflict:
            flash("That email or URL slug is already used by another owner.", "danger")
            conn.close()
            return redirect(url_for("admin_edit_owner", owner_id=owner_id))

        conn.execute("""UPDATE owners SET
            name=?, email=?, phone=?, brand_name=?, slug=?
            WHERE id=?""",
            (name, email, phone, brand_name, slug, owner_id))
        conn.commit()
        conn.close()
        flash(f"Owner '{name}' updated successfully!", "success")
        return redirect(url_for("admin_dashboard"))

    conn.close()
    return render_template("admin_edit_owner.html",
                           owner=dict(owner),
                           admin_name=session["admin_name"])


@app.route("/admin/toggle-owner/<int:owner_id>", methods=["POST"])
@admin_required
@csrf.exempt
def admin_toggle_owner(owner_id):
    conn = get_db()
    owner = conn.execute("SELECT * FROM owners WHERE id=?", (owner_id,)).fetchone()
    if owner:
        new_status = 0 if owner["is_active"] else 1
        conn.execute("UPDATE owners SET is_active=? WHERE id=?", (new_status, owner_id))
        conn.commit()
        flash(f"Owner '{owner['name']}' {'activated' if new_status else 'deactivated'}.", "success")
    conn.close()
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/delete-owner/<int:owner_id>", methods=["POST"])
@admin_required
@csrf.exempt
def admin_delete_owner(owner_id):
    conn = get_db()
    owner = conn.execute("SELECT * FROM owners WHERE id=?", (owner_id,)).fetchone()
    if not owner:
        flash("Owner not found.", "danger")
        conn.close()
        return redirect(url_for("admin_dashboard"))
    name = owner["name"]
    conn.execute("""DELETE FROM enquiries WHERE property_id IN
        (SELECT id FROM properties WHERE owner_id=?)""", (owner_id,))
    conn.execute("DELETE FROM properties WHERE owner_id=?", (owner_id,))
    conn.execute("DELETE FROM owners WHERE id=?", (owner_id,))
    conn.commit()
    conn.close()
    flash(f"Owner '{name}' permanently deleted.", "info")
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/reset-password/<int:owner_id>", methods=["POST"])
@admin_required
@csrf.exempt
def admin_reset_password(owner_id):
    new_password = request.form.get("new_password", "").strip()
    if not new_password:
        flash("Please enter a new password.", "danger")
        return redirect(url_for("admin_view_owner", owner_id=owner_id))
    conn = get_db()
    conn.execute("UPDATE owners SET password=? WHERE id=?",
                 (generate_password_hash(new_password), owner_id))
    conn.commit()
    conn.close()
    flash("Password reset successfully!", "success")
    return redirect(url_for("admin_view_owner", owner_id=owner_id))


@app.route("/admin/settings", methods=["GET", "POST"])
@admin_required
@csrf.exempt
def admin_settings():
    if request.method == "POST":
        username       = request.form.get("mail_username", "").strip()
        password       = request.form.get("mail_password", "").strip()
        sender_name    = request.form.get("mail_sender_name", "").strip()
        admin_whatsapp = request.form.get("admin_whatsapp", "").strip()

        if not username or not sender_name:
            flash("Email address and sender name are required.", "danger")
            return redirect(url_for("admin_settings"))

        save_setting("mail_username",    username)
        save_setting("mail_sender_name", sender_name)
        save_setting("admin_whatsapp",   admin_whatsapp)
        if password:
            save_setting("mail_password", password)

        flash("Settings updated successfully!", "success")
        return redirect(url_for("admin_settings"))

    cfg = get_mail_config()
    admin_whatsapp = get_setting("admin_whatsapp", "")
    return render_template("admin_settings.html",
                           admin_name=session["admin_name"],
                           cfg=cfg,
                           admin_whatsapp=admin_whatsapp)



@app.route("/admin/change-password", methods=["POST"])
@admin_required
@csrf.exempt
def admin_change_password():
    current  = request.form.get("current_password", "").strip()
    new_pw   = request.form.get("new_password", "").strip()
    confirm  = request.form.get("confirm_password", "").strip()

    if not current or not new_pw or not confirm:
        flash("All three fields are required.", "danger")
        return redirect(url_for("admin_settings"))

    if new_pw != confirm:
        flash("New password and confirmation do not match.", "danger")
        return redirect(url_for("admin_settings"))

    if len(new_pw) < 8:
        flash("New password must be at least 8 characters.", "danger")
        return redirect(url_for("admin_settings"))

    # Verify current password
    conn  = get_db()
    admin = conn.execute("SELECT * FROM admins WHERE id=1").fetchone()
    conn.close()

    if not admin or not check_password_hash(admin["password"], current):
        flash("Current password is incorrect.", "danger")
        return redirect(url_for("admin_settings"))

    # Update password
    conn = get_db()
    conn.execute("UPDATE admins SET password=? WHERE id=1",
                 (generate_password_hash(new_pw),))
    conn.commit()
    conn.close()

    flash("Password changed successfully!", "success")
    return redirect(url_for("admin_settings"))


@app.route("/owner-signup", methods=["GET", "POST"])
@csrf.exempt
def owner_signup():
    if request.method == "POST":
        name       = request.form.get("name", "").strip()
        email      = request.form.get("email", "").strip()
        phone      = request.form.get("phone", "").strip()
        profession = request.form.get("profession", "").strip()
        city       = request.form.get("city", "").strip()
        properties = request.form.get("properties", "").strip()
        message    = request.form.get("message", "").strip()

        if not all([name, email, phone]):
            flash("Name, email and phone are required.", "danger")
            return render_template("owner_signup.html")

        # Send email to admin — in background thread so timeout doesn't crash the request
        def send_signup_email():
            try:
                import smtplib, ssl
                from email.mime.multipart import MIMEMultipart
                from email.mime.text import MIMEText
                cfg = get_mail_config()
                html_body = (
                    '<div style="font-family:Arial,sans-serif;max-width:520px;margin:0 auto;">'
                    '<h2 style="color:#7c3aed;">New Owner Signup Request</h2>'
                    '<table style="width:100%;border-collapse:collapse;">'
                    f'<tr><td style="padding:8px 0;color:#6b7280;">Full Name</td><td style="font-weight:700;">{name}</td></tr>'
                    f'<tr><td style="padding:8px 0;color:#6b7280;">Email</td><td style="font-weight:700;">{email}</td></tr>'
                    f'<tr><td style="padding:8px 0;color:#6b7280;">Phone</td><td style="font-weight:700;">{phone}</td></tr>'
                    f'<tr><td style="padding:8px 0;color:#6b7280;">Profession</td><td style="font-weight:700;">{profession or "Not specified"}</td></tr>'
                    f'<tr><td style="padding:8px 0;color:#6b7280;">City</td><td style="font-weight:700;">{city or "Not specified"}</td></tr>'
                    f'<tr><td style="padding:8px 0;color:#6b7280;">Properties</td><td style="font-weight:700;">{properties or "Not specified"}</td></tr>'
                    f'<tr><td style="padding:8px 0;color:#6b7280;">Message</td><td style="font-weight:700;">{message or "None"}</td></tr>'
                    '</table></div>'
                )
                msg = MIMEMultipart("alternative")
                msg["Subject"] = f"New Owner Signup — {name}"
                msg["From"]    = cfg["username"]
                msg["To"]      = cfg["username"]
                msg.attach(MIMEText(html_body, "html"))
                ctx = ssl.create_default_context()
                with smtplib.SMTP_SSL("smtp.hostinger.com", 465, context=ctx, timeout=15) as server:
                    server.login(cfg["username"], cfg["password"])
                    server.sendmail(cfg["username"], cfg["username"], msg.as_string())
                print("Signup email sent OK")
            except Exception as e:
                print(f"Signup email failed (non-fatal): {e}")

        import threading
        threading.Thread(target=send_signup_email, daemon=True).start()

        # Send WhatsApp to admin
        admin_whatsapp = get_setting("admin_whatsapp", "")
        whatsapp_sent = False
        whatsapp_url = ""
        if admin_whatsapp:
            clean = re.sub(r"[^\d]", "", admin_whatsapp)
            wa_text = (
                f"New Owner Signup on ShivyaSpaces!%0A%0A"
                f"Name: {name}%0A"
                f"Email: {email}%0A"
                f"Phone: {phone}%0A"
                f"Profession: {profession or 'N/A'}%0A"
                f"City: {city or 'N/A'}%0A"
                f"Properties: {properties or 'N/A'}%0A"
                f"Message: {message or 'None'}"
            )
            whatsapp_url = f"https://wa.me/{clean}?text={wa_text}"
            whatsapp_sent = True

        flash("Thank you! Your signup request has been submitted. We'll contact you within 24 hours.", "success")
        return render_template("owner_signup.html",
                               submitted=True,
                               whatsapp_sent=whatsapp_sent,
                               whatsapp_url=whatsapp_url if whatsapp_sent else "")

    return render_template("owner_signup.html", submitted=False)


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


@app.route("/termsandconditions")
def termsandconditions():
    return render_template("termsandconditions.html")


@app.route("/data-privacy")
def data_privacy():
    return render_template("data_privacy.html")




@app.route("/tmp-uploads/<filename>")
def serve_tmp_upload(filename):
    """Serve uploaded files from /tmp if static/uploads not writable."""
    from flask import send_from_directory
    import re as _re
    # Security: only allow safe filenames
    if not _re.match(r'^[a-f0-9]{32}\.(jpg|jpeg|png|webp|gif)$', filename):
        from flask import abort
        abort(404)
    return send_from_directory(_tmp_uploads, filename)


@app.route("/upload-image", methods=["POST"])
@csrf.exempt
def upload_image():
    """Upload image — uses Cloudinary if configured, else local static/uploads."""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if not file or file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    # Validate - accept any image type
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else "jpg"
    mime = file.content_type or ""
    if not mime.startswith("image/") and ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": "Please upload an image file (JPG, PNG, WEBP, etc)"}), 400
    if not ext:
        ext = "jpg"

    # Validate size
    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > MAX_UPLOAD_SIZE:
        return jsonify({"error": "Image must be under 5MB"}), 400

    # Try Cloudinary first
    cloud_name = os.environ.get("CLOUDINARY_CLOUD_NAME", "")
    if CLOUDINARY_AVAILABLE and cloud_name:
        try:
            result = cloudinary.uploader.upload(
                file,
                folder="shivyaspaces",
                transformation=[{"width": 1200, "height": 800, "crop": "limit", "quality": "auto"}]
            )
            return jsonify({"url": result["secure_url"]}), 200
        except Exception as e:
            pass  # Fall through to local storage

    # Local storage fallback
    try:
        import uuid
        filename = f"{uuid.uuid4().hex}.{ext}"
        save_path = os.path.join(UPLOAD_FOLDER, filename)
        file.seek(0)
        file.save(save_path)
        url = f"{UPLOAD_URL_BASE}/{filename}"
        return jsonify({"url": url}), 200
    except Exception as e:
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(debug=True)